from io import BytesIO
from pathlib import Path
import base64
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


ORANGE = 'F26426'
DARK_GREY = '5A5A5A'
MARK_PATH = Path('assets/report/spiral-multicolor.png')
CERT_PATH = Path('assets/report/Image21.png')


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=30, left=70, bottom=30, right=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for tag, value in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        node = OxmlElement(f'w:{tag}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def set_table_borders(table, bottom_only=False):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        node = OxmlElement(f'w:{edge}')
        if bottom_only and edge == 'bottom':
            node.set(qn('w:val'), 'single')
            node.set(qn('w:sz'), '10')
            node.set(qn('w:color'), ORANGE)
        else:
            node.set(qn('w:val'), 'nil')
        borders.append(node)
    tbl_pr.append(borders)


def clear_container(container):
    for paragraph in list(container.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)
    for table in list(container.tables):
        table._element.getparent().remove(table._element)


def append_field(paragraph, instruction, fallback='1'):
    field = OxmlElement('w:fldSimple')
    field.set(qn('w:instr'), instruction)
    run = OxmlElement('w:r')
    run_pr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'FFFFFF')
    size = OxmlElement('w:sz')
    size.set(qn('w:val'), '16')
    run_pr.extend([color, size])
    run.append(run_pr)
    text = OxmlElement('w:t')
    text.text = fallback
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def style_run(run, size, color, bold=False):
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    run.bold = bold


def build_header(section):
    header = section.header
    clear_container(header)
    table = header.add_table(rows=1, cols=2, width=Mm(180))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, bottom_only=True)
    left, right = table.rows[0].cells
    left.width = Mm(30)
    right.width = Mm(150)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(left, 10, 20, 10, 20)
    set_cell_margins(right, 10, 20, 10, 20)

    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run().add_picture(str(MARK_PATH), width=Mm(13))

    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('COMPTE RENDU DE VISITE TECHNIQUE')
    style_run(run, 9, (70, 70, 70), bold=True)

    first_header = section.first_page_header
    clear_container(first_header)
    first_header.add_paragraph('')


def build_footer(section):
    footer = section.footer
    clear_container(footer)
    table = footer.add_table(rows=1, cols=2, width=Mm(180))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    left, right = table.rows[0].cells
    left.width = Mm(145)
    right.width = Mm(35)
    set_cell_shading(left, ORANGE)
    set_cell_shading(right, DARK_GREY)
    set_cell_margins(left)
    set_cell_margins(right)

    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('ENERGIE & SERVICE  -  energieetservice.fr')
    style_run(run, 7.5, (255, 255, 255), bold=True)

    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Page ')
    style_run(run, 8, (255, 255, 255))
    append_field(p, 'PAGE')
    run = p.add_run(' / ')
    style_run(run, 8, (255, 255, 255))
    append_field(p, 'NUMPAGES')


def build_first_page_footer(section):
    footer = section.first_page_footer
    clear_container(footer)

    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run('PARIS     NANTES     TOURS     RENNES     BORDEAUX     LYON     CHERBOURG     NÎMES')
    style_run(run, 6.5, (90, 90, 90))

    bar = footer.add_table(rows=1, cols=2, width=Mm(180))
    bar.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(bar)
    left, right = bar.rows[0].cells
    left.width = Mm(148)
    right.width = Mm(32)
    set_cell_shading(left, ORANGE)
    set_cell_shading(right, DARK_GREY)
    set_cell_margins(left, 30, 50, 30, 50)
    set_cell_margins(right, 30, 50, 30, 50)

    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Tél. 01 39 55 17 20  -  143 rue Yves Le Coz - 78000 VERSAILLES  -  contact.versailles@energieetservice.fr')
    style_run(run, 6.1, (255, 255, 255), bold=True)

    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('energieetservice.fr')
    style_run(run, 6.5, (255, 255, 255), bold=True)

    legal = footer.add_table(rows=1, cols=2, width=Mm(180))
    legal.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(legal)
    left, right = legal.rows[0].cells
    left.width = Mm(35)
    right.width = Mm(145)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(left, 10, 20, 10, 20)
    set_cell_margins(right, 10, 20, 10, 20)

    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run().add_picture(str(CERT_PATH), width=Mm(29))

    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('SAS au capital de 292 500 € - Siège social : 143 rue Yves Le Coz - 78000 Versailles - RCS Versailles B 338 335 201 / NAF 7112B')
    style_run(run, 5.5, (120, 120, 120))


def main():
    for asset in (MARK_PATH, CERT_PATH):
        if not asset.is_file() or asset.stat().st_size == 0:
            raise SystemExit(f'Missing Word branding asset: {asset}')

    buffer = BytesIO()
    document = Document()
    normal = document.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(10)

    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    section.header_distance = Mm(6)
    section.footer_distance = Mm(5)
    section.different_first_page_header_footer = True

    paragraph = document.paragraphs[0] if document.paragraphs else document.add_paragraph()
    paragraph.text = 'METRA_TEMPLATE_PLACEHOLDER'

    build_header(section)
    build_footer(section)
    build_first_page_footer(section)
    document.save(buffer)

    payload = buffer.getvalue()
    if not payload.startswith(b'PK'):
        raise SystemExit('Generated DOCX skeleton has no ZIP signature')

    with zipfile.ZipFile(BytesIO(payload), 'r') as archive:
        required = {
            '[Content_Types].xml',
            '_rels/.rels',
            'word/document.xml',
            'word/_rels/document.xml.rels',
            'word/styles.xml',
            'word/settings.xml',
            'word/fontTable.xml',
            'word/theme/theme1.xml',
            'word/header1.xml',
            'word/header2.xml',
            'word/footer1.xml',
            'word/footer2.xml',
            'word/_rels/footer2.xml.rels',
        }
        missing = sorted(required.difference(archive.namelist()))
        if missing:
            raise SystemExit(f'DOCX skeleton incomplete: {missing}')
        document_xml = archive.read('word/document.xml').decode('utf-8')
        if '<w:titlePg' not in document_xml or '<w:headerReference' not in document_xml or '<w:footerReference' not in document_xml:
            raise SystemExit('DOCX skeleton is missing branded first/default header-footer references')

    encoded = base64.b64encode(payload).decode('ascii')
    Path('wordDocxTemplate.generated.js').write_text(
        "// Generated during the Android build from a standards-compliant branded python-docx skeleton.\n"
        f"export const WORD_DOCX_TEMPLATE_BASE64 = '{encoded}';\n",
        encoding='utf-8',
    )
    print(f'Office-compatible branded DOCX skeleton generated: {len(payload)} bytes')


if __name__ == '__main__':
    main()
